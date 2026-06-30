# Programmatic Handbook Generator (DOCX & MD)
import os
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = docx.oxml.parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{fill_hex}"/>')
    tc_pr.append(shd)

def create_docx_handbook():
    doc = Document()
    
    # Page setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Customer Churn Prediction & LTV Engine")
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(26, 54, 93) # Deep Navy
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Detailed Project Handbook & Code Explanations (Weeks 1 & 2)")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(74, 85, 104) # Grey
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("\n" + "="*60 + "\n")
    
    # Introduction Section
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Project Overview & Team Structure")
    h1_run.font.color.rgb = RGBColor(26, 54, 93)
    
    p = doc.add_paragraph()
    p.add_run(
        "This handbook details the codebase, implementation pipeline, and individual responsibilities "
        "for the Customer Churn Prediction project. The project is structured across a 5-member team "
        "to deliver a robust Data Engineering, Data Analysis, and Machine Learning workflow. "
        "This volume focuses on Setup, Week 1 (Data Engineering & EDA), and Week 2 (Churn Modeling & SHAP Explainability)."
    )
    
    # Table of responsibilities
    doc.add_heading("Team Allocations - Weeks 1 & 2", level=2)
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Shading Accent 1'
    
    headers = ["Member Role", "Week 1 Responsibilities", "Week 2 Responsibilities"]
    col_widths = [Inches(1.8), Inches(2.4), Inches(2.4)]
    
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], "1A365D")
        
    members_data = [
        ("Member 1: Data Engineer", 
         "PostgreSQL Schema Setup; db_connection.py creation; etl_pipeline.py ingestion script.", 
         "Refactoring preprocessing; support model preprocessors (preprocess.py)."),
        ("Member 2: Data Analyst", 
         "Jupyter EDA Notebook (EDA.ipynb); Business Insights Report (business_insights.md).", 
         "Evaluate model outputs; verify F1 and ROC AUC metrics; write presenter explanations."),
        ("Member 3: ML Engineer", 
         "Verify raw data structure; align schemas for machine learning input variables.", 
         "Model pipeline creation; model training (Logistic Regression, Random Forest, XGBoost); select best classifier."),
        ("Member 4: SHAP Explainer", 
         "Pre-modeling feature analysis; identify potential collinearity in correlation heatmap.", 
         "Calculate SHAP explainability values; generate global summary plots (shap_explain.py)."),
        ("Member 5: BI & Testing", 
         "Setup dashboard data templates; outline database loading validation metrics.", 
         "Generate target prediction outputs; export churn_predictions.csv; prepare dashboard inputs.")
    ]
    
    for row_idx, data in enumerate(members_data):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            # Width formatting
            row_cells[col_idx].width = col_widths[col_idx]
            
    doc.add_paragraph("\n")
    
    # Week 1 Section
    h1_2 = doc.add_heading(level=1)
    h1_2_run = h1_2.add_run("2. Week 1 Code & Component Explanations")
    h1_2_run.font.color.rgb = RGBColor(26, 54, 93)
    
    # File: create_tables.sql
    doc.add_heading("A. SQL Schema Setup (create_tables.sql & load_data.sql)", level=2)
    doc.add_paragraph(
        "Responsible: Member 1 (Data Engineer)\n"
        "Description: Sets up raw and cleaned relational structures in the PostgreSQL database."
    )
    p_code = doc.add_paragraph()
    p_code_run = p_code.add_run(
        "Code Logic:\n"
        "- Creates the 'telco_customers' table with text/numeric types matching the CSV format. TotalCharges is imported as VARCHAR to allow blank space ingestion.\n"
        "- Creates the 'cleaned_telco_customers' table where Churn is mapped to INTEGER (0/1) and TotalCharges is parsed into NUMERIC.\n"
        "- Specifies the 'COPY' SQL command inside load_data.sql to ingest the local dataset."
    )
    p_code_run.font.italic = True
    
    # File: db_connection.py
    doc.add_heading("B. Database Connector Setup (db_connection.py)", level=2)
    doc.add_paragraph(
        "Responsible: Member 1 (Data Engineer)\n"
        "Description: Establishes a database engine using SQLAlchemy."
    )
    p_db = doc.add_paragraph()
    p_db_run = p_db.add_run(
        "Code Logic:\n"
        "- Attempts to build a PostgreSQL connection engine using variables: DB_USER, DB_PASSWORD, DB_HOST, DB_PORT, and DB_NAME.\n"
        "- Gracefully falls back to sqlite:///customer_churn_db.sqlite if connection fails (e.g. database offline during local dev)."
    )
    p_db_run.font.italic = True
    
    # File: etl_pipeline.py
    doc.add_heading("C. Data Ingestion Pipeline (etl_pipeline.py)", level=2)
    doc.add_paragraph(
        "Responsible: Member 1 (Data Engineer)\n"
        "Description: Implements Extract, Transform, and Load (ETL) mechanics."
    )
    p_etl = doc.add_paragraph()
    p_etl_run = p_etl.add_run(
        "Code Logic:\n"
        "1. Extract: Loads raw CSV data into a Pandas DataFrame.\n"
        "2. Load Raw: Saves raw records to table 'telco_customers'.\n"
        "3. Transform: Coerces TotalCharges to float, fills blanks with median, drops non-predictive customerID, maps Churn (No/Yes) to (0/1).\n"
        "4. Load Cleaned: Saves processed rows to 'cleaned_telco_customers' and writes data/cleaned_telco.csv."
    )
    p_etl_run.font.italic = True
    
    # File: EDA.ipynb
    doc.add_heading("D. Exploratory Data Analysis (EDA.ipynb)", level=2)
    doc.add_paragraph(
        "Responsible: Member 2 (Data Analyst)\n"
        "Description: Runs statistical inspections and generates visualization charts."
    )
    p_eda = doc.add_paragraph()
    p_eda_run = p_eda.add_run(
        "Code Logic:\n"
        "- Target Check: Measures Churn class distribution (26.54% Churn vs 73.46% Retained).\n"
        "- Numeric Distributions: Plots histplots for tenure (identifies onboarding spike) and monthly charges.\n"
        "- Relationship Bars: Computes churn rate across Contract Type (month-to-month contracts have 42.71% churn) and Payment Method (Electronic check has 45.29% churn).\n"
        "- Heatmap: Correlates numeric attributes, showing that tenure and Churn are negatively correlated (-0.35)."
    )
    p_eda_run.font.italic = True
    
    # File: business_insights.md
    doc.add_heading("E. Presentation Insights (business_insights.md)", level=2)
    doc.add_paragraph(
        "Responsible: Member 2 (Data Analyst)\n"
        "Description: Summarizes insights and presentation speaking scripts for each graph."
    )
    p_ins = doc.add_paragraph()
    p_ins_run = p_ins.add_run(
        "Contents:\n"
        "- Highlights major retention blockers (onboarding stage churn, high bills).\n"
        "- Speaking scripts: Provides specific speaking points for the PPT slides to explain distributions, contract barriers, and payment friction.\n"
        "- Actions: Recommends incentives for migrating users to auto-pay and long-term contracts."
    )
    p_ins_run.font.italic = True
    
    doc.add_paragraph("\n")
    
    # Week 2 Section
    h1_3 = doc.add_heading(level=1)
    h1_3_run = h1_3.add_run("3. Week 2 Code & Component Explanations")
    h1_3_run.font.color.rgb = RGBColor(26, 54, 93)
    
    # File: preprocess.py
    doc.add_heading("A. Preprocessing Refactoring (preprocess.py)", level=2)
    doc.add_paragraph(
        "Responsible: Member 1 (Data Cleaning & Preprocessing)\n"
        "Description: Contains the standalone cleaning workflow to resolve file path dependencies."
    )
    p_prep = doc.add_paragraph()
    p_prep_run = p_prep.add_run(
        "Code Logic:\n"
        "- Refactored the hardcoded local absolute paths into relative paths.\n"
        "- Converts TotalCharges to numeric, handles blank fields, drops customerID, maps Churn, and outputs data/cleaned_telco.csv."
    )
    p_prep_run.font.italic = True
    
    # File: evaluate_models.py
    doc.add_heading("B. Scoring Interface Setup (evaluate_models.py)", level=2)
    doc.add_paragraph(
        "Responsible: Member 2 (Model Evaluation)\n"
        "Description: Utility scoring code for consistent metrics computation."
    )
    p_ev = doc.add_paragraph()
    p_ev_run = p_ev.add_run(
        "Code Logic:\n"
        "- Resolves the Git conflicts present in the initial templates.\n"
        "- Defines an 'evaluate()' function calculating: Accuracy, Precision, Recall, F1-Score, and ROC-AUC score.\n"
        "- Returns the scores as a dictionary to facilitate model comparison."
    )
    p_ev_run.font.italic = True
    
    # File: train_churn.py
    doc.add_heading("C. ML Pipeline Training (train_churn.py)", level=2)
    doc.add_paragraph(
        "Responsible: Member 3 (ML Churn Model Development)\n"
        "Description: Trains and compares 3 models to select the best predictor."
    )
    p_trn = doc.add_paragraph()
    p_trn_run = p_trn.add_run(
        "Code Logic:\n"
        "1. Feature Engineering: Computes AvgChargePerMonth = TotalCharges / (tenure + 1).\n"
        "2. Pipeline Design: Implements a Scikit-Learn Pipeline wrapping a ColumnTransformer (OneHotEncoder for object columns, StandardScaler for numeric columns) and the Classifier. This prevents data leakage during training.\n"
        "3. Splitting: Stratifies the train-test split (80/20) to preserve target ratios.\n"
        "4. Inits: Trains Logistic Regression, Random Forest, and XGBoost.\n"
        "5. Selection: Evaluates test metrics. Logistic Regression is selected as the best overall classifier (F1: 0.6039, ROC AUC: 0.8473) and saved as models/churn_model.pkl.\n"
        "6. Prediction Export: Saves data/churn_predictions.csv containing probabilities for dashboard usage."
    )
    p_trn_run.font.italic = True
    
    # File: shap_explain.py
    doc.add_heading("D. SHAP Explainability Script (shap_explain.py)", level=2)
    doc.add_paragraph(
        "Responsible: Member 4 (SHAP Explainability)\n"
        "Description: Quantifies individual feature impact on model decisions."
    )
    p_shp = doc.add_paragraph()
    p_shp_run = p_shp.add_run(
        "Code Logic:\n"
        "- Loads models/churn_model.pkl and the cleaned dataset.\n"
        "- Extracts features and transforms them using the pipeline's preprocessor.\n"
        "- Obtains post-encoding feature names (e.g. Contract_One year, PaymentMethod_Electronic check).\n"
        "- Instantiates shap.Explainer, calculates SHAP values, and saves a summary plot to reports/shap_summary_plot.png."
    )
    p_shp_run.font.italic = True
    
    # File: churn_predictions.csv
    doc.add_heading("E. Predictions Dataset Generation (churn_predictions.csv)", level=2)
    doc.add_paragraph(
        "Responsible: Member 5 (Testing & Dashboard Dataset)\n"
        "Description: Exports the predictions dataset containing probabilities."
    )
    p_prd = doc.add_paragraph()
    p_prd_run = p_prd.add_run(
        "Code Logic:\n"
        "- Created during the execution of train_churn.py.\n"
        "- Concatenates test set features with ActualChurn, PredictedChurn, and ChurnProbability.\n"
        "- Used directly as input for BI tools (Metabase, PowerBI) to analyze customer churn risk levels."
    )
    p_prd_run.font.italic = True
    
    doc.add_paragraph("\n" + "="*60 + "\n")
    
    # Save document
    os.makedirs("reports", exist_ok=True)
    doc_path = "reports/Project_Handbook_Weeks_1_and_2.docx"
    doc.save(doc_path)
    print(f"Word Document saved successfully at {doc_path}!")

def create_markdown_handbook():
    md_content = """# Customer Churn Prediction & LTV Engine
## Detailed Project Handbook & Code Explanations (Weeks 1 & 2)

---

## 1. Project Overview & Team Structure

This handbook details the codebase, implementation pipeline, and individual responsibilities for the Customer Churn Prediction project. The project is structured across a 5-member team to deliver a robust Data Engineering, Data Analysis, and Machine Learning workflow. This volume focuses on Setup, Week 1 (Data Engineering & EDA), and Week 2 (Churn Modeling & SHAP Explainability).

### Team Allocations - Weeks 1 & 2

| Member Role | Week 1 Responsibilities | Week 2 Responsibilities |
| :--- | :--- | :--- |
| **Member 1: Data Engineer** | PostgreSQL Schema Setup; `db_connection.py` creation; `etl_pipeline.py` ingestion script. | Refactoring preprocessing; support model preprocessors (`preprocess.py`). |
| **Member 2: Data Analyst** | Jupyter EDA Notebook (`EDA.ipynb`); Business Insights Report (`business_insights.md`). | Evaluate model outputs; verify F1 and ROC AUC metrics; write presenter explanations. |
| **Member 3: ML Engineer** | Verify raw data structure; align schemas for machine learning input variables. | Model pipeline creation; model training (Logistic Regression, Random Forest, XGBoost); select best classifier. |
| **Member 4: SHAP Explainer** | Pre-modeling feature analysis; identify potential collinearity in correlation heatmap. | Calculate SHAP explainability values; generate global summary plots (`shap_explain.py`). |
| **Member 5: BI & Testing** | Setup dashboard data templates; outline database loading validation metrics. | Generate target prediction outputs; export `churn_predictions.csv`; prepare dashboard inputs. |

---

## 2. Week 1 Code & Component Explanations

### A. SQL Schema Setup (`sql/create_tables.sql` & `sql/load_data.sql`)
*   **Responsible**: Member 1 (Data Engineer)
*   **Description**: Sets up raw and cleaned relational structures in the PostgreSQL database.
*   **Code Logic**:
    *   Creates the `telco_customers` table with text/numeric types matching the CSV format. `TotalCharges` is imported as `VARCHAR` to allow blank space ingestion.
    *   Creates the `cleaned_telco_customers` table where `Churn` is mapped to `INTEGER` (0/1) and `TotalCharges` is parsed into `NUMERIC`.
    *   Specifies the `COPY` SQL command inside `load_data.sql` to ingest the local dataset.

### B. Database Connector Setup (`scripts/db_connection.py`)
*   **Responsible**: Member 1 (Data Engineer)
*   **Description**: Establishes a database engine using SQLAlchemy.
*   **Code Logic**:
    *   Attempts to build a PostgreSQL connection engine using variables: `DB_USER`, `DB_PASSWORD`, `DB_HOST`, `DB_PORT`, and `DB_NAME`.
    *   Gracefully falls back to `sqlite:///customer_churn_db.sqlite` if connection fails (e.g. database offline during local dev).

### C. Data Ingestion Pipeline (`scripts/etl_pipeline.py`)
*   **Responsible**: Member 1 (Data Engineer)
*   **Description**: Implements Extract, Transform, and Load (ETL) mechanics.
*   **Code Logic**:
    1.  **Extract**: Loads raw CSV data into a Pandas DataFrame.
    2.  **Load Raw**: Saves raw records to table `telco_customers`.
    3.  **Transform**: Coerces `TotalCharges` to float, fills blanks with median, drops non-predictive `customerID`, maps `Churn` (`No`/`Yes`) to (`0`/`1`).
    4.  **Load Cleaned**: Saves processed rows to `cleaned_telco_customers` and writes `data/cleaned_telco.csv`.

### D. Exploratory Data Analysis (`notebooks/EDA.ipynb`)
*   **Responsible**: Member 2 (Data Analyst)
*   **Description**: Runs statistical inspections and generates visualization charts.
*   **Code Logic**:
    *   **Target Check**: Measures Churn class distribution (26.54% Churn vs 73.46% Retained).
    *   **Numeric Distributions**: Plots histplots for `tenure` (identifies onboarding spike) and monthly charges.
    *   **Relationship Bars**: Computes churn rate across Contract Type (month-to-month contracts have 42.71% churn) and Payment Method (Electronic check has 45.29% churn).
    *   **Heatmap**: Correlates numeric attributes, showing that tenure and Churn are negatively correlated (-0.35).

### E. Presentation Insights (`reports/business_insights.md`)
*   **Responsible**: Member 2 (Data Analyst)
*   **Description**: Summarizes insights and presentation speaking scripts for each graph.
*   **Contents**:
    *   Highlights major retention blockers (onboarding stage churn, high bills).
    *   **Speaking scripts**: Provides specific speaking points for the PPT slides to explain distributions, contract barriers, and payment friction.
    *   **Actions**: Recommends incentives for migrating users to auto-pay and long-term contracts.

---

## 3. Week 2 Code & Component Explanations

### A. Preprocessing Refactoring (`scripts/preprocess.py`)
*   **Responsible**: Member 1 (Data Cleaning & Preprocessing)
*   **Description**: Contains the standalone cleaning workflow to resolve file path dependencies.
*   **Code Logic**:
    *   Refactored the hardcoded local absolute paths into relative paths.
    *   Converts `TotalCharges` to numeric, handles blank fields, drops `customerID`, maps `Churn`, and outputs `data/cleaned_telco.csv`.

### B. Scoring Interface Setup (`scripts/evaluate_models.py`)
*   **Responsible**: Member 2 (Model Evaluation)
*   **Description**: Utility scoring code for consistent metrics computation.
*   **Code Logic**:
    *   Resolves the Git conflicts present in the initial templates.
    *   Defines an `evaluate()` function calculating: Accuracy, Precision, Recall, F1-Score, and ROC-AUC score.
    *   Returns the scores as a dictionary to facilitate model comparison.

### C. ML Pipeline Training (`scripts/train_churn.py`)
*   **Responsible**: Member 3 (ML Churn Model Development)
*   **Description**: Trains and compares 3 models to select the best predictor.
*   **Code Logic**:
    1.  **Feature Engineering**: Computes `AvgChargePerMonth` = `TotalCharges` / (`tenure` + 1).
    2.  **Pipeline Design**: Implements a Scikit-Learn Pipeline wrapping a `ColumnTransformer` (`OneHotEncoder` for object columns, `StandardScaler` for numeric columns) and the Classifier. This prevents data leakage during training.
    3.  **Splitting**: Stratifies the train-test split (80/20) to preserve target ratios.
    4.  **Inits**: Trains Logistic Regression, Random Forest, and XGBoost.
    5.  **Selection**: Evaluates test metrics. Logistic Regression is selected as the best overall classifier (F1: 0.6039, ROC AUC: 0.8473) and saved as `models/churn_model.pkl`.
    6.  **Prediction Export**: Saves `data/churn_predictions.csv` containing probabilities for dashboard usage.

### D. SHAP Explainability Script (`scripts/shap_explain.py`)
*   **Responsible**: Member 4 (SHAP Explainability)
*   **Description**: Quantifies individual feature impact on model decisions.
*   **Code Logic**:
    *   Loads `models/churn_model.pkl` and the cleaned dataset.
    *   Extracts features and transforms them using the pipeline's preprocessor.
    *   Obtains post-encoding feature names (e.g. `Contract_One year`, `PaymentMethod_Electronic check`).
    *   Instantiates `shap.Explainer`, calculates SHAP values, and saves a summary plot to `reports/shap_summary_plot.png`.

### E. Predictions Dataset Generation (`data/churn_predictions.csv`)
*   **Responsible**: Member 5 (Testing & Dashboard Dataset)
*   **Description**: Exports the predictions dataset containing probabilities.
*   **Code Logic**:
    *   Created during the execution of `train_churn.py`.
    *   Concatenates test set features with `ActualChurn`, `PredictedChurn`, and `ChurnProbability`.
    *   Used directly as input for BI tools (Metabase, PowerBI) to analyze customer churn risk levels.
"""
    
    md_path = "reports/Project_Handbook_Weeks_1_and_2.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"Markdown Handbook saved successfully at {md_path}!")

if __name__ == "__main__":
    create_docx_handbook()
    create_markdown_handbook()
